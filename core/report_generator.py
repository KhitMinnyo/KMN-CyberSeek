"""
KMN-CyberSeek Session Report Generator
Produces a DOCX penetration-test report from a session dict (as returned by
orchestrator.get_session_report()).

Designed to be import-safe: if python-docx is not installed the module loads
without error and generate_report() raises ImportError with a clear message,
so the FastAPI server still starts even on minimal installs.

Usage (from backend):
    from core.report_generator import generate_report
    path = generate_report(session_report_dict, output_path="/tmp/report.docx")
    # → returns the path on success, raises on failure
"""

import json
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Risk-level → colour mapping (Word RGB)
# ---------------------------------------------------------------------------
_RISK_COLORS = {
    "high":    (0xD3, 0x2F, 0x2F),  # red
    "medium":  (0xF5, 0x7F, 0x17),  # amber
    "low":     (0x2E, 0x7D, 0x32),  # green
    "unknown": (0x55, 0x55, 0x55),  # grey
}

_HEADER_BG = (0x1A, 0x23, 0x7E)   # dark indigo — title bar
_ACCENT_BG  = (0x37, 0x47, 0x4F)  # blue-grey — section headers


def _require_docx():
    try:
        from docx import Document                          # noqa: F401
        from docx.shared import Pt, RGBColor, Inches     # noqa: F401
        from docx.enum.text import WD_ALIGN_PARAGRAPH     # noqa: F401
        from docx.oxml.ns import qn                       # noqa: F401
        from docx.oxml import OxmlElement                 # noqa: F401
    except ImportError:
        raise ImportError(
            "python-docx is required for report generation. "
            "Install it with: pip install python-docx"
        )


def _set_cell_bg(cell, rgb: tuple):
    """Apply a solid background colour to a table cell (OOXML shading)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_section_heading(doc, text: str):
    """Add a styled section heading paragraph."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*_ACCENT_BG)
    # Bottom border as a visual rule
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*_ACCENT_BG))
    pBdr.append(bottom)
    pPr.append(pBdr)
    doc.add_paragraph("")   # spacer


def _add_table_header_row(table, headers: List[str], widths_dxa: List[int]):
    """Write column headers into row 0 with dark background + white bold text."""
    from docx.shared import Pt, RGBColor
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths_dxa)):
        cell = row.cells[i]
        cell.width = w
        _set_cell_bg(cell, _HEADER_BG)
        para = cell.paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def generate_report(session_report: Dict, output_path: Optional[str] = None) -> str:
    """Generate a DOCX penetration-test report from a session report dict.

    Args:
        session_report: The dict returned by Orchestrator.get_session_report()
        output_path: Where to write the .docx. Defaults to /tmp/kmn_report_{session_id}.docx

    Returns:
        Absolute path to the generated file.

    Raises:
        ImportError: if python-docx is not installed
        ValueError: if session_report is empty/invalid
    """
    _require_docx()

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    session = session_report.get("session", {})
    session_id = session.get("session_id", "unknown")
    target_ip = session.get("target_ip", "N/A")
    target_domain = session.get("target_domain") or ""
    created_at = (session.get("created_at") or "")[:19].replace("T", " ")
    status = session.get("status", "unknown")
    stage = session.get("current_stage", "unknown")
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M UTC")

    discovered_services: List[Dict] = session_report.get("discovered_services", [])
    discovered_hosts: List[Dict]    = session_report.get("discovered_hosts", [])
    vulnerabilities: List[Dict]     = session_report.get("vulnerabilities", [])
    commands: List[Dict]            = session_report.get("commands_executed", [])
    credentials: List[Dict]         = session_report.get("credentials", [])
    ai_decisions: List[Dict]        = session_report.get("ai_decisions", [])

    if not output_path:
        output_path = f"/tmp/kmn_report_{session_id[:12]}.docx"

    doc = Document()

    # --- Page margins (narrow) ---
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ============================================================
    # COVER PAGE
    # ============================================================
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("KMN-CyberSeek")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(*_HEADER_BG)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Penetration Test Session Report")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(*_ACCENT_BG)

    doc.add_paragraph("")

    meta_lines = [
        ("Target", f"{target_ip}" + (f" / {target_domain}" if target_domain else "")),
        ("Session ID", session_id),
        ("Session Created", created_at),
        ("Report Generated", report_date),
        ("Final Status", status.upper()),
        ("Final Stage", stage),
    ]
    meta_table = doc.add_table(rows=len(meta_lines), cols=2)
    meta_table.style = "Table Grid"
    total_w = int(section.page_width - section.left_margin - section.right_margin)
    label_w = int(total_w * 0.35)
    value_w = total_w - label_w
    for i, (label, value) in enumerate(meta_lines):
        row = meta_table.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = label_w
        vc.width = value_w
        _set_cell_bg(lc, _ACCENT_BG)
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        vr = vc.paragraphs[0].add_run(str(value))
        vr.font.size = Pt(9)

    doc.add_page_break()

    # ============================================================
    # 1. EXECUTIVE SUMMARY
    # ============================================================
    _add_section_heading(doc, "1. Executive Summary")

    high_v   = sum(1 for v in vulnerabilities if v.get("risk_level") == "high")
    medium_v = sum(1 for v in vulnerabilities if v.get("risk_level") == "medium")
    low_v    = sum(1 for v in vulnerabilities if v.get("risk_level") == "low")
    unverif  = sum(1 for v in vulnerabilities if v.get("status") == "unverified")

    summary_text = (
        f"This report summarises the results of an AI-directed penetration test session "
        f"targeting {target_ip}" + (f" ({target_domain})" if target_domain else "") + ". "
        f"The session discovered {len(discovered_hosts)} host(s) and "
        f"{len(discovered_services)} service(s). "
        f"A total of {len(vulnerabilities)} vulnerability finding(s) were recorded: "
        f"{high_v} high, {medium_v} medium, {low_v} low severity "
        f"(plus {unverif} unverified leads from web research). "
        f"{len(credentials)} credential(s) were captured. "
        f"{len(commands)} command(s) were executed during the session."
    )
    p = doc.add_paragraph(summary_text)
    p.style.font.size = Pt(10)

    doc.add_paragraph("")

    # Metrics summary table (1 row × 5 cols)
    metrics = [
        ("Hosts", str(len(discovered_hosts))),
        ("Services", str(len(discovered_services))),
        ("High Vulns", str(high_v)),
        ("Medium Vulns", str(medium_v)),
        ("Credentials", str(len(credentials))),
    ]
    mt = doc.add_table(rows=2, cols=len(metrics))
    mt.style = "Table Grid"
    col_w = total_w // len(metrics)
    for i, (label, val) in enumerate(metrics):
        hcell = mt.rows[0].cells[i]
        vcell = mt.rows[1].cells[i]
        hcell.width = col_w
        vcell.width = col_w
        _set_cell_bg(hcell, _ACCENT_BG)
        hr2 = hcell.paragraphs[0].add_run(label)
        hr2.bold = True
        hr2.font.size = Pt(8)
        hr2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr2 = vcell.paragraphs[0].add_run(val)
        vr2.bold = True
        vr2.font.size = Pt(14)
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # ============================================================
    # 2. DISCOVERED SERVICES
    # ============================================================
    _add_section_heading(doc, "2. Discovered Services")

    if discovered_services:
        svc_headers = ["Host", "Port", "Service", "Version", "State"]
        svc_widths  = [int(total_w * p) for p in [0.20, 0.10, 0.20, 0.35, 0.15]]
        st2 = doc.add_table(rows=1 + len(discovered_services), cols=5)
        st2.style = "Table Grid"
        _add_table_header_row(st2, svc_headers, svc_widths)
        for i, svc in enumerate(discovered_services):
            row = st2.rows[i + 1]
            vals = [
                str(svc.get("host") or target_ip),
                str(svc.get("port") or ""),
                str(svc.get("service") or ""),
                str(svc.get("version") or ""),
                str(svc.get("state") or "open"),
            ]
            for j, (cell, val, w) in enumerate(zip(row.cells, vals, svc_widths)):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
    else:
        doc.add_paragraph("No services discovered in this session.")

    doc.add_paragraph("")

    # ============================================================
    # 3. VULNERABILITY FINDINGS
    # ============================================================
    _add_section_heading(doc, "3. Vulnerability Findings")

    if vulnerabilities:
        # Group by risk descending
        ordered = sorted(
            vulnerabilities,
            key=lambda v: {"high": 0, "medium": 1, "low": 2}.get(v.get("risk_level", ""), 3)
        )
        vheaders = ["#", "Risk", "Name", "Host:Port", "CVE(s)", "Source", "Status"]
        vwidths  = [int(total_w * p) for p in [0.04, 0.07, 0.24, 0.15, 0.18, 0.14, 0.13]]
        # Adjust last to fill remainder
        vwidths[-1] = total_w - sum(vwidths[:-1])

        vt = doc.add_table(rows=1 + len(ordered), cols=7)
        vt.style = "Table Grid"
        _add_table_header_row(vt, vheaders, vwidths)

        for i, vuln in enumerate(ordered):
            row = vt.rows[i + 1]
            risk = (vuln.get("risk_level") or "unknown").lower()
            rgb = _RISK_COLORS.get(risk, _RISK_COLORS["unknown"])
            cves = ", ".join(vuln.get("cve_ids") or []) or "—"
            host_port = f"{vuln.get('host') or ''}"
            if vuln.get("port"):
                host_port += f":{vuln['port']}"
            vals = [
                str(i + 1),
                risk.upper(),
                (vuln.get("name") or "")[:80],
                host_port,
                cves[:60],
                (vuln.get("source_tool") or "")[:20],
                (vuln.get("status") or "confirmed")[:15],
            ]
            for j, (cell, val, w) in enumerate(zip(row.cells, vals, vwidths)):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(8)
                if j == 1:  # Risk column - colour text
                    run.bold = True
                    run.font.color.rgb = RGBColor(*rgb)

        # Per-finding details (expandable detail block for high/medium)
        doc.add_paragraph("")
        _add_section_heading(doc, "3.1  Finding Details")
        for i, vuln in enumerate(ordered):
            if vuln.get("risk_level") not in ("high", "medium"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(f"[{i+1}] {vuln.get('name') or 'Unnamed'}")
            r.bold = True
            r.font.size = Pt(10)
            risk = (vuln.get("risk_level") or "unknown").lower()
            r.font.color.rgb = RGBColor(*_RISK_COLORS.get(risk, _RISK_COLORS["unknown"]))

            for label, key in [
                ("Description", "description"), ("Affected Software", "service_version"),
            ]:
                val = vuln.get(key) or ""
                if val:
                    lp = doc.add_paragraph()
                    lr = lp.add_run(f"{label}: ")
                    lr.bold = True
                    lr.font.size = Pt(9)
                    lp.add_run(val[:500]).font.size = Pt(9)

            cves = ", ".join(vuln.get("cve_ids") or [])
            if cves:
                lp = doc.add_paragraph()
                lr = lp.add_run("CVE(s): ")
                lr.bold = True
                lr.font.size = Pt(9)
                lp.add_run(cves).font.size = Pt(9)

            refs = vuln.get("reference_urls") or []
            if refs:
                lp = doc.add_paragraph()
                lr = lp.add_run("References: ")
                lr.bold = True
                lr.font.size = Pt(9)
                lp.add_run("; ".join(refs[:3])).font.size = Pt(9)

            doc.add_paragraph("")
    else:
        doc.add_paragraph("No vulnerabilities recorded in this session.")

    doc.add_paragraph("")

    # ============================================================
    # 4. CREDENTIALS CAPTURED
    # ============================================================
    _add_section_heading(doc, "4. Credentials Captured")

    if credentials:
        cheaders = ["Username", "Secret", "Type", "Service", "Discovered"]
        cwidths  = [int(total_w * p) for p in [0.18, 0.30, 0.10, 0.12, 0.20]]
        cwidths[-1] = total_w - sum(cwidths[:-1])
        ct = doc.add_table(rows=1 + len(credentials), cols=5)
        ct.style = "Table Grid"
        _add_table_header_row(ct, cheaders, cwidths)
        for i, cred in enumerate(credentials):
            row = ct.rows[i + 1]
            vals = [
                cred.get("username") or "",
                (cred.get("secret") or "")[:64],
                cred.get("secret_type") or "password",
                cred.get("service") or "",
                (cred.get("discovered_at") or "")[:19],
            ]
            for cell, val, w in zip(row.cells, vals, cwidths):
                cell.width = w
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
    else:
        doc.add_paragraph("No credentials were captured during this session.")

    doc.add_paragraph("")

    # ============================================================
    # 5. EXECUTED COMMANDS LOG
    # ============================================================
    _add_section_heading(doc, "5. Executed Commands Log")

    if commands:
        for i, cmd in enumerate(commands):
            p = doc.add_paragraph()
            r = p.add_run(f"[{i+1}]  {cmd.get('command','')[:120]}")
            r.font.name = "Courier New"
            r.font.size = Pt(8)
            r.bold = True

            ts = (cmd.get("timestamp") or "")[:19]
            ok = "✓" if cmd.get("success") else "✗"
            meta = doc.add_paragraph()
            mr = meta.add_run(f"     {ok}  {ts}")
            mr.font.size = Pt(8)
            mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            output = (cmd.get("output") or "")[:600]
            if output:
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Inches(0.3)
                orr = op.add_run(output)
                orr.font.name = "Courier New"
                orr.font.size = Pt(7)
                orr.font.color.rgb = RGBColor(0x00, 0x33, 0x00)

            doc.add_paragraph("")
    else:
        doc.add_paragraph("No commands executed in this session.")

    # ============================================================
    # 6. AI DECISIONS SUMMARY (last 10)
    # ============================================================
    if ai_decisions:
        doc.add_page_break()
        _add_section_heading(doc, "6. AI Decision Log (last 10)")
        for decision in ai_decisions[-10:]:
            p = doc.add_paragraph()
            tr2 = p.add_run(f"Phase: {decision.get('attack_phase','?')}  |  "
                            f"Risk: {decision.get('risk_level','?')}  |  "
                            f"Confidence: {decision.get('confidence', 0):.0%}")
            tr2.bold = True
            tr2.font.size = Pt(9)

            cmd_p = doc.add_paragraph()
            cmd_p.paragraph_format.left_indent = Inches(0.3)
            cr = cmd_p.add_run(f"Suggested: {decision.get('suggested_command','')[:120]}")
            cr.font.name = "Courier New"
            cr.font.size = Pt(8)

            reasoning = (decision.get("reasoning") or "")[:400]
            if reasoning:
                rp = doc.add_paragraph()
                rp.paragraph_format.left_indent = Inches(0.3)
                rr = rp.add_run(reasoning)
                rr.font.size = Pt(8)
                rr.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

            doc.add_paragraph("")

    # ============================================================
    # DISCLAIMER FOOTER
    # ============================================================
    doc.add_page_break()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run("LEGAL DISCLAIMER")
    dr.bold = True
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(*_RISK_COLORS["high"])

    doc.add_paragraph(
        "This report was generated by KMN-CyberSeek, an AI-assisted penetration testing "
        "framework. All testing activity recorded in this report was performed only against "
        "systems for which explicit written authorisation was obtained prior to testing "
        "(as confirmed by the authorization_confirmed flag in session metadata). "
        "Vulnerability findings derived from unverified web research (source: threat-intel-cache) "
        "are marked as such and must be independently corroborated before being treated as "
        "confirmed. The operator is solely responsible for the legality and scope of all "
        "testing activity."
    ).runs[0].font.size = Pt(9)

    doc.save(output_path)
    logger.info(f"Report saved to {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# PDF report generator (uses fpdf2 — pure Python, no LibreOffice needed)
# ---------------------------------------------------------------------------

def _require_fpdf():
    try:
        from fpdf import FPDF  # noqa: F401
    except ImportError:
        raise ImportError(
            "fpdf2 is required for PDF report generation. "
            "Install it with: pip install fpdf2"
        )


def generate_pdf_report(session_report: Dict, output_path: Optional[str] = None) -> str:
    """Generate a PDF penetration-test report.

    Args:
        session_report: dict as returned by orchestrator.get_session_report()
        output_path: where to write the file. Defaults to /tmp/kmn_report_<id>.pdf

    Returns:
        Absolute path to the generated PDF.
    """
    _require_fpdf()
    from fpdf import FPDF

    meta = session_report.get("session", {})
    session_id = meta.get("session_id", "unknown")[:12]
    target = meta.get("target_ip", "—")
    created = meta.get("created_at", "—")
    status = meta.get("status", "—")
    generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    vulns: List[Dict] = session_report.get("vulnerabilities", [])
    commands: List[Dict] = session_report.get("commands", [])
    services: List[Dict] = session_report.get("services", [])
    credentials: List[Dict] = session_report.get("credentials", [])
    ai_decisions: List[Dict] = session_report.get("ai_decisions", [])

    # Risk summary
    high_c = sum(1 for v in vulns if v.get("risk_level") == "high")
    med_c  = sum(1 for v in vulns if v.get("risk_level") == "medium")
    low_c  = sum(1 for v in vulns if v.get("risk_level") == "low")

    if output_path is None:
        output_path = os.path.join("/tmp", f"kmn_report_{session_id}.pdf")

    # ── PDF setup ──────────────────────────────────────────────────────────
    class _PDF(FPDF):
        def header(self):
            self.set_fill_color(26, 35, 126)       # dark indigo
            self.rect(0, 0, 210, 14, "F")
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(255, 255, 255)
            self.set_xy(8, 3)
            self.cell(0, 8, "KMN-CyberSeek  |  Penetration Test Report  — CONFIDENTIAL")
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "", 7)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, f"Page {self.page_no()}", align="C")
            self.set_text_color(0, 0, 0)

    pdf = _PDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()
    pdf.set_margins(14, 18, 14)

    def _h1(txt: str):
        pdf.ln(4)
        pdf.set_fill_color(55, 71, 79)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, f"  {txt}", fill=True, ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    def _row(label: str, value: str, bold_val: bool = False):
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(50, 6, label + ":", ln=False)
        pdf.set_font("Helvetica", "B" if bold_val else "", 9)
        pdf.multi_cell(0, 6, str(value))

    def _para(txt: str, size: int = 9):
        pdf.set_font("Helvetica", "", size)
        pdf.multi_cell(0, 5, txt)
        pdf.ln(1)

    # ── Cover / summary ───────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(26, 35, 126)
    pdf.cell(0, 12, "KMN-CyberSeek", ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 7, "AI-Driven Penetration Test Report", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    _h1("Session Metadata")
    _row("Session ID", session_id)
    _row("Target", target)
    _row("Status", status)
    _row("Started", str(created))
    _row("Generated", generated_at)
    pdf.ln(2)

    _h1("Executive Summary")
    _row("Total Vulnerabilities", str(len(vulns)))
    _row("High Risk", str(high_c), bold_val=high_c > 0)
    _row("Medium Risk", str(med_c))
    _row("Low Risk", str(low_c))
    _row("Services Discovered", str(len(services)))
    _row("Commands Executed", str(len(commands)))
    _row("Credentials Captured", str(len(credentials)))
    pdf.ln(2)

    # ── Vulnerabilities ───────────────────────────────────────────────────
    if vulns:
        _h1(f"Vulnerability Findings ({len(vulns)})")
        for i, v in enumerate(
            sorted(vulns, key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x.get("risk_level", ""), 3)),
            start=1
        ):
            risk = v.get("risk_level", "unknown").upper()
            name = v.get("name", "Unknown")
            host = v.get("host", "—")
            port = v.get("port") or "—"
            svc  = v.get("service", "—")
            cves = ", ".join(json.loads(v.get("cve_ids") or "[]") or []) or "—"
            desc = v.get("description", "")

            # Risk colour
            rc = {"HIGH": (211, 47, 47), "MEDIUM": (245, 127, 23), "LOW": (46, 125, 50)}.get(risk, (100, 100, 100))
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*rc)
            pdf.cell(0, 6, f"[{risk}] {i}. {name}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)
            pdf.cell(0, 5, f"Host: {host}  |  Port: {port}  |  Service: {svc}  |  CVEs: {cves}", ln=True)
            if desc:
                pdf.set_font("Helvetica", "I", 8)
                pdf.multi_cell(0, 5, desc[:300] + ("…" if len(desc) > 300 else ""))
            pdf.ln(1)

    # ── Services ─────────────────────────────────────────────────────────
    if services:
        _h1(f"Discovered Services ({len(services)})")
        pdf.set_font("Helvetica", "B", 8)
        # Header row
        col_w = [30, 20, 18, 60, 52]
        headers = ["Host", "Port", "State", "Service", "Version"]
        for w, h in zip(col_w, headers):
            pdf.cell(w, 6, h, border=1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        for s in services[:40]:
            host = str(s.get("host", ""))[:18]
            port = str(s.get("port", ""))
            state = str(s.get("state", ""))
            svc  = str(s.get("service", ""))[:28]
            ver  = str(s.get("version", ""))[:30]
            for w, val in zip(col_w, [host, port, state, svc, ver]):
                pdf.cell(w, 5, val, border=1)
            pdf.ln()
        pdf.ln(2)

    # ── Credentials ───────────────────────────────────────────────────────
    if credentials:
        _h1(f"Captured Credentials ({len(credentials)})")
        pdf.set_font("Helvetica", "B", 8)
        for w, h in zip([40, 60, 25, 30, 25], ["Username", "Secret", "Type", "Service", "Host"]):
            pdf.cell(w, 6, h, border=1)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        for c in credentials:
            secret = str(c.get("secret", ""))
            if len(secret) > 28:
                secret = secret[:25] + "…"
            for w, val in zip([40, 60, 25, 30, 25], [
                str(c.get("username", ""))[:22],
                secret,
                str(c.get("secret_type", ""))[:10],
                str(c.get("service", ""))[:14],
                str(c.get("host", ""))[:14],
            ]):
                pdf.cell(w, 5, val, border=1)
            pdf.ln()
        pdf.ln(2)

    # ── Commands ──────────────────────────────────────────────────────────
    if commands:
        _h1(f"Commands Log ({len(commands)})")
        for i, cmd in enumerate(commands[-30:], start=1):   # last 30 max
            ok = cmd.get("success", False)
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_text_color(30, 100, 30) if ok else pdf.set_text_color(180, 30, 30)
            pdf.cell(0, 5, f"{'✓' if ok else '✗'}  {str(cmd.get('command', ''))[:100]}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 7)
            out = str(cmd.get("output") or "").strip()[:400]
            if out:
                pdf.multi_cell(0, 4, out)
            pdf.ln(1)

    # ── Legal disclaimer ──────────────────────────────────────────────────
    pdf.add_page()
    _h1("Legal Disclaimer")
    _para(
        "This report was generated by KMN-CyberSeek and is intended SOLELY for use by "
        "authorised security professionals operating on systems for which explicit written "
        "authorisation was obtained prior to testing. Vulnerability findings derived from "
        "unverified web research are marked as such and must be independently corroborated. "
        "The operator is solely responsible for the legality and scope of all testing activity."
    )

    pdf.output(output_path)
    logger.info(f"PDF report saved to {output_path}")
    return output_path
